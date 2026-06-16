"""Pure bytes-to-text extractors for Drive file content.

One function per supported mime type. Each takes raw bytes and returns
``(text, method_label)`` where method_label matches what gets written
into ``DocumentText.extraction_method``.

Heavy parser libraries (pymupdf, python-docx, openpyxl) are imported
lazily inside each function so this module loads even when the
``[content]`` optional dependency group is not installed -- callers that
hit an unsupported mime are routed to ``skip`` paths in
``content_pipeline.py`` before they get here.

If a parser is genuinely missing for a mime type we *do* handle, the
function returns ``(None, 'failed-no-parser')`` rather than raising, so
the pipeline keeps going.
"""

from __future__ import annotations

import logging
from collections.abc import Callable

logger = logging.getLogger(__name__)

# Mime types this module knows how to extract.  The pipeline checks this
# set first; anything not listed gets a 'skipped-mime' DocumentText row
# without bytes being downloaded.
SUPPORTED_MIMES: set[str] = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # .xlsx
    "application/vnd.google-apps.document",
    "application/vnd.google-apps.spreadsheet",
}

# Mime types served via Drive's export endpoint (Google-native formats).
# Pipeline must call ``export_google_doc`` not ``download_file`` for these.
GOOGLE_NATIVE_MIMES: dict[str, str] = {
    "application/vnd.google-apps.document": "text/plain",
    "application/vnd.google-apps.spreadsheet": "text/csv",
}


# ---------------------------------------------------------------------------
# Individual extractors
# ---------------------------------------------------------------------------


def extract_pdf(raw: bytes) -> tuple[str | None, str]:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("pymupdf not installed; install with `pip install .[content]`")
        return None, "failed-no-parser"

    try:
        doc = fitz.open(stream=raw, filetype="pdf")
    except Exception as exc:
        logger.warning("PDF parse failed: %s", exc)
        return None, "failed-parse"

    try:
        parts = [page.get_text("text") for page in doc]
    finally:
        doc.close()

    text = "\n".join(parts).strip()
    return (text or None), "pdf-pymupdf"


def extract_docx(raw: bytes) -> tuple[str | None, str]:
    """Extract text from a .docx Word document."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("python-docx not installed; install with `pip install .[content]`")
        return None, "failed-no-parser"

    import io

    try:
        doc = DocxDocument(io.BytesIO(raw))
    except Exception as exc:
        logger.warning("DOCX parse failed: %s", exc)
        return None, "failed-parse"

    # Paragraph text first, then table cells (often important in contracts).
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    text = "\n".join(parts).strip()
    return (text or None), "docx-python"


# Spreadsheets are the worst offenders for unbounded extraction: a single
# acquisition model produced 2.18M chars of text, and feeding a wall of bare
# numbers to the LLM yields garbage "amounts" with no direction.  We therefore
# (a) keep each sheet's HEADER row visible so the model can associate values
# with columns, (b) cap rows + chars PER SHEET, and (c) cap the whole workbook,
# noting any truncation so the model knows content continues.
_XLSX_MAX_ROWS_PER_SHEET = 200
_XLSX_MAX_CHARS_PER_SHEET = 16_000
_XLSX_MAX_TOTAL_CHARS = 60_000


# Trailing empty columns are common in exported sheets; drop them so rows aren't
# a forest of tabs.
def _trim_trailing_empty(cells: list) -> list:
    end = len(cells)
    while end > 0 and (cells[end - 1] is None or str(cells[end - 1]).strip() == ""):
        end -= 1
    return cells[:end]


def extract_xlsx(raw: bytes) -> tuple[str | None, str]:
    """Extract text from a .xlsx Excel workbook, structure-preserving + bounded.

    Each sheet is rendered as TSV with its header row kept visible, capped to a
    sane number of rows/chars (a workbook is not a contract -- an uncapped
    acquisition model is megabytes of noise).  ``data_only=True`` so formula
    cells return their last-cached computed value; sheets that were never
    recalculated come back mostly blank and are flagged.
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        logger.warning("openpyxl not installed; install with `pip install .[content]`")
        return None, "failed-no-parser"

    import io

    try:
        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    except Exception as exc:
        logger.warning("XLSX parse failed: %s", exc)
        return None, "failed-parse"

    parts: list[str] = []
    total_chars = 0
    try:
        for sheet in wb.worksheets:
            if total_chars >= _XLSX_MAX_TOTAL_CHARS:
                parts.append("### (further sheets omitted -- workbook too large)")
                break
            sheet_lines: list[str] = [f"### {sheet.title}"]
            sheet_chars = 0
            rows_emitted = 0
            rows_seen = 0
            nonempty = 0
            for row in sheet.iter_rows(values_only=True):
                rows_seen += 1
                cells = _trim_trailing_empty(list(row))
                if not any(c not in (None, "") for c in cells):
                    continue
                nonempty += 1
                if (
                    rows_emitted >= _XLSX_MAX_ROWS_PER_SHEET
                    or sheet_chars >= _XLSX_MAX_CHARS_PER_SHEET
                ):
                    continue  # keep counting rows_seen so the note is accurate
                line = "\t".join("" if c is None else str(c) for c in cells)
                sheet_lines.append(line)
                sheet_chars += len(line) + 1
                rows_emitted += 1
            if nonempty == 0:
                # Almost always a never-recalculated formula sheet (data_only
                # returns None) -- say so rather than emit an empty section.
                sheet_lines.append("(no values -- sheet may contain only uncomputed formulas)")
            elif nonempty > rows_emitted:
                sheet_lines.append(
                    f"(... {nonempty - rows_emitted} more row(s) not shown; "
                    f"sheet truncated for length)"
                )
            block = "\n".join(sheet_lines)
            parts.append(block)
            total_chars += len(block) + 1
    finally:
        wb.close()
    text = "\n".join(parts).strip()
    return (text or None), "xlsx-openpyxl"


def extract_gdoc_export(raw: bytes) -> tuple[str | None, str]:
    """Decode a Google Doc exported as text/plain.

    Drive's export endpoint returns UTF-8 bytes directly -- no parsing
    needed, just decode safely.
    """
    try:
        text = raw.decode("utf-8", errors="replace").strip()
    except Exception as exc:
        logger.warning("Google Doc decode failed: %s", exc)
        return None, "failed-parse"
    return (text or None), "gdoc-export"


def extract_gsheet_export(raw: bytes) -> tuple[str | None, str]:
    """Decode a Google Sheet exported as text/csv."""
    try:
        text = raw.decode("utf-8", errors="replace").strip()
    except Exception as exc:
        logger.warning("Google Sheet decode failed: %s", exc)
        return None, "failed-parse"
    return (text or None), "gsheet-export"


# Mime -> extractor function map.  Pipeline uses this dispatch table.
_EXTRACTORS: dict[str, Callable[[bytes], tuple[str | None, str]]] = {
    "application/pdf": extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_docx,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": extract_xlsx,
    "application/vnd.google-apps.document": extract_gdoc_export,
    "application/vnd.google-apps.spreadsheet": extract_gsheet_export,
}


def extract_text(raw: bytes, mime_type: str) -> tuple[str | None, str]:
    """Dispatch to the right extractor for *mime_type*.

    Returns ``(None, 'skipped-mime')`` if the mime is unsupported -- callers
    should normally check ``SUPPORTED_MIMES`` first to avoid wasted bytes.
    """
    fn = _EXTRACTORS.get(mime_type)
    if fn is None:
        return None, "skipped-mime"
    return fn(raw)


# ---------------------------------------------------------------------------
# Token-count heuristic
# ---------------------------------------------------------------------------


def estimate_tokens(text: str | None) -> int | None:
    """Rough token count without a real tokenizer dep.

    Claude/GPT tokenizers average ~3.5-4 chars/token for English prose;
    we use 4 as a conservative floor (slightly under-counts long words,
    slightly over-counts code).  Good enough for context-window budgeting.
    Swap to tiktoken when accuracy matters.
    """
    if not text:
        return None
    return max(1, len(text) // 4)
